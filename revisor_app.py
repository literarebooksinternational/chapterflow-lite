# revisor_app.py

import io
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import docx
from docx.shared import RGBColor
from spellchecker import SpellChecker
import re

# Inicializa o aplicativo Flask e o CORS para permitir requisições do seu front-end
app = Flask(__name__)
CORS(app)

# --- Dicionários de Correção ---
# Dicionários com as substituições a serem feitas.
# É fácil adicionar novas regras aqui no futuro.

# Cacoetes e vícios de linguagem
import re

# CACOETES: substituições de tiques de fala ou palavras repetitivas
CACOETES = {
    re.compile(r'\b(Aí|Ai|Aih) eu fui\b', re.IGNORECASE): 'Então eu fui',
    re.compile(r'\b(Aí|Ai|Aih) eu disse\b', re.IGNORECASE): 'Então eu disse',
    re.compile(r'\btipo assim\b', re.IGNORECASE): 'por exemplo',
    re.compile(r'\bné\b', re.IGNORECASE): '',  # Remove "né" soltos
    re.compile(r'\b(okay|ok)\b', re.IGNORECASE): 'certo',
    re.compile(r'\b(Então|Tipo|Aí)\b', re.IGNORECASE): '',  # Remove muletas
    re.compile(r'\b(entendeu\?|sacou\?)\b', re.IGNORECASE): '',  # Retira perguntas muleta
    re.compile(r'\b(Ééé+|Hum+|Ah+)\b', re.IGNORECASE): '',  # Sons prolongados
    re.compile(r'\b(no fundo|na real|sabe)\b', re.IGNORECASE): '',  # Expressões repetitivas
    re.compile(r'\b(assim tipo|coisa|negócio)\b', re.IGNORECASE): 'por exemplo',  # Substitui palavras vagas
    re.compile(r'\b(então assim)\b', re.IGNORECASE): 'portanto',
    re.compile(r'\b(tá bom|beleza|massa)\b', re.IGNORECASE): 'certo',
}

# GIRIAS E TERMOS INFORMAIS
GIRIAS = {
    re.compile(r'\bnois\b', re.IGNORECASE): 'nós',
    re.compile(r'\btamo\b', re.IGNORECASE): 'estamos',
    re.compile(r'\bta\b', re.IGNORECASE): 'está',
    re.compile(r'\bvc\b', re.IGNORECASE): 'você',
    re.compile(r'\bpra\b', re.IGNORECASE): 'para',
    re.compile(r'\bpros?\b', re.IGNORECASE): 'para os',
    re.compile(r'\bnéh?\b', re.IGNORECASE): '',  # Remove "néh" desnecessário
    re.compile(r'\bblz\b', re.IGNORECASE): 'certo',
    re.compile(r'\bkkk+\b', re.IGNORECASE): '',  # Remove risadas
    re.compile(r'\bpq\b', re.IGNORECASE): 'porque',
    re.compile(r'\bq\b', re.IGNORECASE): 'que',
    re.compile(r'\bvcê\b', re.IGNORECASE): 'você',
    re.compile(r'\bmané\b', re.IGNORECASE): 'cara',
    re.compile(r'\bmano\b', re.IGNORECASE): 'amigo',
    re.compile(r'\btbm\b', re.IGNORECASE): 'também',
    re.compile(r'\bai+\b', re.IGNORECASE): 'ah',
}


# Configura o corretor ortográfico para o português
spell = SpellChecker(language='pt')

def aplicar_correcoes_e_marcar(doc):
    """
    Função principal que processa o documento Word.
    Itera por cada parágrafo, aplica as correções e realça as alterações.

    Args:
        doc (docx.Document): O objeto do documento Word carregado.

    Returns:
        tuple: Uma tupla contendo:
            - doc (docx.Document): O documento modificado.
            - report (list): Uma lista de dicionários com os detalhes de cada alteração.
    """
    report = []
    
    # Combina todos os dicionários de regras para verificação
    todas_regras = {**CACOETES, **GIRIAS}

    for p in doc.paragraphs:
        # Pula parágrafos vazios
        if not p.text.strip():
            continue

        texto_original_paragrafo = p.text
        texto_corrigido = texto_original_paragrafo

        # 1. Aplica as regras de gírias e cacoetes
        for pattern, replacement in todas_regras.items():
            # Usa uma função para registrar a alteração enquanto substitui
            def registrar_e_substituir(match):
                original = match.group(0)
                # Verifica se a substituição não é a mesma para evitar loops/falsos positivos
                if original.lower() != replacement.lower():
                    report.append({'original': original, 'corrigido': replacement, 'tipo': 'Regra'})
                return replacement
            
            texto_corrigido, count = pattern.subn(registrar_e_substituir, texto_corrigido)
        
        # 2. Aplica a correção ortográfica
        palavras = re.findall(r'\b\w+\b', texto_corrigido)
        misspelled = spell.unknown(palavras)
        
        texto_final_ortografia = texto_corrigido
        for palavra in misspelled:
            # Não corrige palavras que começam com maiúscula (nomes próprios, etc.)
            if not palavra[0].isupper():
                correcao = spell.correction(palavra)
                if correcao and correcao != palavra:
                    # Usa regex para substituir a palavra inteira e evitar substituir substrings
                    pattern_palavra = re.compile(r'\b' + re.escape(palavra) + r'\b')
                    texto_final_ortografia = pattern_palavra.sub(correcao, texto_final_ortografia)
                    report.append({'original': palavra, 'corrigido': correcao, 'tipo': 'Ortografia'})
        
        # 3. Reconstrói o parágrafo com o realce se houver mudanças
        if texto_original_paragrafo != texto_final_ortografia:
            # Limpa o parágrafo original
            p.clear()
            
            # Adiciona o texto corrigido, mas agora realçando
            # Esta é a nossa alternativa ao "Track Changes"
            run = p.add_run(texto_final_ortografia)
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            # Mantém a formatação original (negrito, itálico) se possível
            # (Simplificado aqui, mas a formatação principal é mantida no nível do parágrafo)

    return doc, report


@app.route('/revisar-documento', methods=['POST'])
def revisar_documento():
    """
    Endpoint da API que recebe o arquivo .docx.
    """
    if 'file' not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nome de arquivo vazio"}), 400

    if file and file.filename.endswith('.docx'):
        try:
            # Lê o arquivo em memória
            doc = docx.Document(file)
            
            # Aplica as correções
            doc_corrigido, relatorio = aplicar_correcoes_e_marcar(doc)
            
            # Salva o documento corrigido em um buffer de bytes
            buffer = io.BytesIO()
            doc_corrigido.save(buffer)
            buffer.seek(0)
            
            # Prepara a resposta
            # Precisamos enviar o arquivo e o relatório JSON.
            # O ideal é usar uma resposta multipart, mas para simplificar,
            # vamos fazer duas requisições no front-end ou usar um truque.
            # Aqui, vamos retornar o arquivo e o front-end pedirá o relatório depois.
            # Melhor ainda: retornaremos o relatório nos headers, se for pequeno.
            # A melhor solução: retornar o arquivo como file e o json no mesmo response.
            
            response = send_file(
                buffer,
                as_attachment=True,
                download_name=f"revisado_{file.filename}",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            # Adiciona o relatório como um header customizado (convertido para JSON string)
            # Headers têm limite de tamanho, mas para um relatório de texto é geralmente ok.
            import json
            response.headers['X-Corrections-Report'] = json.dumps(relatorio)
            # Permite que o front-end acesse este header customizado
            response.headers['Access-Control-Expose-Headers'] = 'X-Corrections-Report'

            return response

        except Exception as e:
            return jsonify({"error": f"Erro ao processar o arquivo: {str(e)}"}), 500

    return jsonify({"error": "Formato de arquivo inválido. Envie um .docx"}), 400

if __name__ == '__main__':
    # Roda o servidor na porta 5000 por padrão
    app.run(debug=True, port=5000)