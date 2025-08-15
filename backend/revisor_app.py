# revisor_app.py

import io
import re
import json
import docx
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from spellchecker import SpellChecker
from docx.shared import RGBColor

# Importa a nova biblioteca para comparar textos
import diff_match_patch as dmp_module

# Inicializa o aplicativo Flask e o CORS
app = Flask(__name__)
CORS(app)

# --- Dicionários de Correção Expandidos ---

# Ortografia / Acentuação
ERRORS = {
    re.compile(r'\bpor favor\b', re.IGNORECASE): 'por favor',
    re.compile(r'\bcaxia\b', re.IGNORECASE): 'caixa',
    re.compile(r'\bexeção\b', re.IGNORECASE): 'exceção',
    re.compile(r'\bconheser\b', re.IGNORECASE): 'conhecer',
    re.compile(r'\bseção\b', re.IGNORECASE): 'seção',
    re.compile(r'\bdefeito\b', re.IGNORECASE): 'defeito',
    re.compile(r'\bprincípio\b', re.IGNORECASE): 'princípio',
    re.compile(r'\binteressante\b', re.IGNORECASE): 'interessante',
    re.compile(r'\bparalelo\b', re.IGNORECASE): 'paralelo',
    re.compile(r'\bprivilégio\b', re.IGNORECASE): 'privilégio',
    re.compile(r'\bbenefício\b', re.IGNORECASE): 'benefício',
    re.compile(r'\bbeneficiente\b', re.IGNORECASE): 'beneficente',
    re.compile(r'\bmim\s+mesmo\b', re.IGNORECASE): 'eu mesmo',
    re.compile(r'\benxergar\b', re.IGNORECASE): 'enxergar',
    re.compile(r'\bmexer\b', re.IGNORECASE): 'mexer',
    re.compile(r'\bbruxa\b', re.IGNORECASE): 'bruxa',
    re.compile(r'\bfaxina\b', re.IGNORECASE): 'faxina',
    re.compile(r'\bpeixe\b', re.IGNORECASE): 'peixe',
    re.compile(r'\bchuva\b', re.IGNORECASE): 'chuva',
    re.compile(r'\bxadrez\b', re.IGNORECASE): 'xadrez',
    re.compile(r'\bxícara\b', re.IGNORECASE): 'xícara',
    re.compile(r'\bexame\b', re.IGNORECASE): 'exame',
    re.compile(r'\bexemplo\b', re.IGNORECASE): 'exemplo',
    re.compile(r'\bexercício\b', re.IGNORECASE): 'exercício',
    re.compile(r'\bexperiência\b', re.IGNORECASE): 'experiência',
    re.compile(r'\bexplicação\b', re.IGNORECASE): 'explicação',
    re.compile(r'\bextremamente\b', re.IGNORECASE): 'extremamente',
    re.compile(r'\bmáximo\b', re.IGNORECASE): 'máximo',
    re.compile(r'\bpróximo\b', re.IGNORECASE): 'próximo',
    re.compile(r'\btóxico\b', re.IGNORECASE): 'tóxico',
    re.compile(r'\bsintaxe\b', re.IGNORECASE): 'sintaxe',
    re.compile(r'\banálize\b', re.IGNORECASE): 'análise',
    re.compile(r'\bpêsames\b', re.IGNORECASE): 'pêsames',
    re.compile(r'\bênfase\b', re.IGNORECASE): 'ênfase',
    re.compile(r'\btênis\b', re.IGNORECASE): 'tênis',
    re.compile(r'\bgênesis\b', re.IGNORECASE): 'gênesis',
    re.compile(r'\bhipótese\b', re.IGNORECASE): 'hipótese',
    re.compile(r'\bparentêsis\b', re.IGNORECASE): 'parênteses',
    re.compile(r'\bciúme\b', re.IGNORECASE): 'ciúme',
    re.compile(r'\bviúva\b', re.IGNORECASE): 'viúva',
    re.compile(r'\babraço\b', re.IGNORECASE): 'abraço',
    re.compile(r'\bergueu\b', re.IGNORECASE): 'ergueu',
    re.compile(r'\btábua\b', re.IGNORECASE): 'tábua',
    re.compile(r'\btáboas\b', re.IGNORECASE): 'tábuas',
    re.compile(r'\bmágoa\b', re.IGNORECASE): 'mágoa',
    re.compile(r'\blíngua\b', re.IGNORECASE): 'língua',
    re.compile(r'\bágua\b', re.IGNORECASE): 'água',
}

# Concordância
CONCORD = {
    re.compile(r'\bnós vai(s?)(\b|$)', re.IGNORECASE): 'nós vamos',
    re.compile(r'\bcujo o\b', re.IGNORECASE): 'cujo',
    re.compile(r'\bà uma\b', re.IGNORECASE): 'a uma',
    re.compile(r'\bà partir\b', re.IGNORECASE): 'a partir',
    re.compile(r'\bà respeito\b', re.IGNORECASE): 'a respeito',
    re.compile(r'\bà toa\b', re.IGNORECASE): 'à toa',
    re.compile(r'\bhaviam muitas pessoas\b', re.IGNORECASE): 'havia muitas pessoas',
    re.compile(r'\bexistem pessoas\b', re.IGNORECASE): 'existem pessoas',
    re.compile(r'\bfazem dois anos\b', re.IGNORECASE): 'faz dois anos',
    re.compile(r'\bhouveram problemas\b', re.IGNORECASE): 'houve problemas',
    re.compile(r'\bbastante pessoas\b', re.IGNORECASE): 'bastantes pessoas',
    re.compile(r'\bmenas pessoas\b', re.IGNORECASE): 'menos pessoas',
    re.compile(r'\balerta os usuários\b', re.IGNORECASE): 'alerta aos usuários',
    re.compile(r'\bincluso as mulheres\b', re.IGNORECASE): 'inclusive as mulheres',
    re.compile(r'\bobrigado\b(?=.*\bela\b)', re.IGNORECASE): 'obrigada',
    re.compile(r'\bmesmo ela\b', re.IGNORECASE): 'mesmo ela',
    re.compile(r'\beles mesmo\b', re.IGNORECASE): 'eles mesmos',
    re.compile(r'\belas mesmo\b', re.IGNORECASE): 'elas mesmas',
    re.compile(r'\banexo segue\b', re.IGNORECASE): 'anexo seguem' ,
    re.compile(r'\bé proibido entrada\b', re.IGNORECASE): 'é proibida a entrada',
    re.compile(r'\bé necessário paciência\b', re.IGNORECASE): 'é necessária paciência',
    re.compile(r'\bmeia xícara\b', re.IGNORECASE): 'meia xícara',
    re.compile(r'\bmeio nervosa\b', re.IGNORECASE): 'meio nervosa',
    re.compile(r'\bos óculos está\b', re.IGNORECASE): 'os óculos estão',
    re.compile(r'\ba maioria foram\b', re.IGNORECASE): 'a maioria foi',
}

# Homônimos
HOM = {
    re.compile(r'\bconcerto\b(?=.*reparo|.*arrumar|.*consertar)', re.IGNORECASE): 'conserto',
    re.compile(r'\bconserto\b(?=.*música|.*orquestra|.*sinfonia)', re.IGNORECASE): 'concerto',
    re.compile(r'\bseção\b(?=.*cortar|.*dividir)', re.IGNORECASE): 'secção',
    re.compile(r'\bsessão\b(?=.*cinema|.*reunião|.*tempo)', re.IGNORECASE): 'sessão',
    re.compile(r'\bcessão\b(?=.*ceder|.*dar)', re.IGNORECASE): 'cessão',
    re.compile(r'\btráfego\b(?=.*trânsito|.*veículos)', re.IGNORECASE): 'tráfego',
    re.compile(r'\btráfico\b(?=.*drogas|.*ilegal)', re.IGNORECASE): 'tráfico',
    re.compile(r'\bemerge\b(?=.*sair|.*aparecer)', re.IGNORECASE): 'emerge',
    re.compile(r'\bimerge\b(?=.*mergulhar|.*entrar)', re.IGNORECASE): 'imerge',
    re.compile(r'\binfligir\b(?=.*aplicar|.*impor)', re.IGNORECASE): 'infligir',
    re.compile(r'\binfringir\b(?=.*violar|.*desrespeitar)', re.IGNORECASE): 'infringir',
    re.compile(r'\bflagrante\b(?=.*evidente|.*óbvio)', re.IGNORECASE): 'flagrante',
    re.compile(r'\bfragrante\b(?=.*cheiro|.*perfume)', re.IGNORECASE): 'fragrante',
    re.compile(r'\bemergir\b(?=.*sair)', re.IGNORECASE): 'emergir',
    re.compile(r'\bimergir\b(?=.*mergulhar)', re.IGNORECASE): 'imergir',
    re.compile(r'\bdiscriminar\b(?=.*separar|.*diferenciar)', re.IGNORECASE): 'discriminar',
    re.compile(r'\bdiscriminar\b(?=.*preconceito)', re.IGNORECASE): 'discriminar',
    re.compile(r'\bespiar\b(?=.*olhar)', re.IGNORECASE): 'espiar',
    re.compile(r'\bexpiar\b(?=.*pagar|.*culpa)', re.IGNORECASE): 'expiar',
}

# Preposições
PREP = {
    re.compile(r'\bgostar (?!de)', re.IGNORECASE): 'gostar de',
    re.compile(r'\bassisti (?!a)', re.IGNORECASE): 'assistir a',
    re.compile(r'\bchegar (?!a)', re.IGNORECASE): 'chegar a',
    re.compile(r'\bvisar (?!a)', re.IGNORECASE): 'visar a',
    re.compile(r'\baspira (?!a)', re.IGNORECASE): 'aspirar a',
    re.compile(r'\bobedece (?!a)', re.IGNORECASE): 'obedecer a',
    re.compile(r'\bresiste (?!a)', re.IGNORECASE): 'resistir a',
    re.compile(r'\bprefere (?!.*a.*)', re.IGNORECASE): 'preferir... a...',
    re.compile(r'\bprefer.*mais.*do que\b', re.IGNORECASE): 'preferir... a...',
    re.compile(r'\bnamora (?!com)', re.IGNORECASE): 'namorar com',
    re.compile(r'\bpisa (?!em)', re.IGNORECASE): 'pisar em',
    re.compile(r'\bsimpática (?!com)', re.IGNORECASE): 'simpática com',
    re.compile(r'\bdiferente (?!de)', re.IGNORECASE): 'diferente de',
    re.compile(r'\bigual (?!a)', re.IGNORECASE): 'igual a',
    re.compile(r'\bpreferível (?!a)', re.IGNORECASE): 'preferível a',
    re.compile(r'\bsuperior (?!a)', re.IGNORECASE): 'superior a',
    re.compile(r'\binferior (?!a)', re.IGNORECASE): 'inferior a',
    re.compile(r'\banterior (?!a)', re.IGNORECASE): 'anterior a',
    re.compile(r'\bposterior (?!a)', re.IGNORECASE): 'posterior a',
    re.compile(r'\bproibir (?!de)', re.IGNORECASE): 'proibir de',
    re.compile(r'\bdesistir (?!de)', re.IGNORECASE): 'desistir de',
    re.compile(r'\besquecer (?!de)', re.IGNORECASE): 'esquecer de',
    re.compile(r'\blembrar (?!de)', re.IGNORECASE): 'lembrar de',
}

# Expressões redundantes / estilo informal
STYLE = {
    re.compile(r'\bIrei fazer\b', re.IGNORECASE): 'Farei',
    re.compile(r'\bvou estar fazendo\b', re.IGNORECASE): 'vou fazer',
    re.compile(r'\bestou a fazer\b', re.IGNORECASE): 'estou fazendo',
    re.compile(r'\bsubir para cima\b', re.IGNORECASE): 'subir',
    re.compile(r'\bdescer para baixo\b', re.IGNORECASE): 'descer',
    re.compile(r'\bentrar para dentro\b', re.IGNORECASE): 'entrar',
    re.compile(r'\bsair para fora\b', re.IGNORECASE): 'sair',
    re.compile(r'\bvoltar de novo\b', re.IGNORECASE): 'voltar',
    re.compile(r'\brepetir de novo\b', re.IGNORECASE): 'repetir',
    re.compile(r'\bcriação nova\b', re.IGNORECASE): 'criação',
    re.compile(r'\binovação nova\b', re.IGNORECASE): 'inovação',
    re.compile(r'\breabertura nova\b', re.IGNORECASE): 'reabertura',
    re.compile(r'\bemonopólio exclusivo\b', re.IGNORECASE): 'monopólio',
    re.compile(r'\bconsenso geral\b', re.IGNORECASE): 'consenso',
    re.compile(r'\belo de ligação\b', re.IGNORECASE): 'elo',
    re.compile(r'\bfato real\b', re.IGNORECASE): 'fato',
    re.compile(r'\bplanos futuros\b', re.IGNORECASE): 'planos',
    re.compile(r'\berário público\b', re.IGNORECASE): 'erário',
    re.compile(r'\bhistória do passado\b', re.IGNORECASE): 'história',
    re.compile(r'\bmeta a atingir\b', re.IGNORECASE): 'meta',
    re.compile(r'\bpopulação do povo\b', re.IGNORECASE): 'população',
    re.compile(r'\bunir junto\b', re.IGNORECASE): 'unir',
    re.compile(r'\bjuntar junto\b', re.IGNORECASE): 'juntar',
    re.compile(r'\bcerteza absoluta\b', re.IGNORECASE): 'certeza',
    re.compile(r'\bcadáver morto\b', re.IGNORECASE): 'cadáver',
    re.compile(r'\bhemorragia de sangue\b', re.IGNORECASE): 'hemorragia',
    re.compile(r'\bpequenos detalhes\b', re.IGNORECASE): 'detalhes',
    re.compile(r'\bviúva do falecido\b', re.IGNORECASE): 'viúva',
    re.compile(r'\bconviver junto\b', re.IGNORECASE): 'conviver',
    re.compile(r'\bencarar de frente\b', re.IGNORECASE): 'encarar',
    re.compile(r'\bcompilar junto\b', re.IGNORECASE): 'compilar',
}

# Ortografia informal
INFORMAL = {
    re.compile(r'\boque\b', re.IGNORECASE): 'o que',
    re.compile(r'\bq\b', re.IGNORECASE): 'que',
    re.compile(r'\btb\b', re.IGNORECASE): 'também',
    re.compile(r'\btbm\b', re.IGNORECASE): 'também',
    re.compile(r'\bvc\b', re.IGNORECASE): 'você',
    re.compile(r'\bvcs\b', re.IGNORECASE): 'vocês',
    re.compile(r'\bpq\b', re.IGNORECASE): 'porque',
    re.compile(r'\bpor que\b(?=\s*\?)', re.IGNORECASE): 'por que',
    re.compile(r'\bporque\b(?=.*\?)', re.IGNORECASE): 'por que',
    re.compile(r'\bporque\b(?=.*[.!])', re.IGNORECASE): 'porque',
    re.compile(r'\bpra\b', re.IGNORECASE): 'para',
    re.compile(r'\bpro\b', re.IGNORECASE): 'para o',
    re.compile(r'\bpros\b', re.IGNORECASE): 'para os',
    re.compile(r'\bpras\b', re.IGNORECASE): 'para as',
    re.compile(r'\bneh\b', re.IGNORECASE): '',
    re.compile(r'\bne\b', re.IGNORECASE): '',
    re.compile(r'\btá\b', re.IGNORECASE): 'está',
    re.compile(r'\btô\b', re.IGNORECASE): 'estou',
    re.compile(r'\btamo\b', re.IGNORECASE): 'estamos',
    re.compile(r'\btão\b', re.IGNORECASE): 'estão',
    re.compile(r'\bvamo\b', re.IGNORECASE): 'vamos',
    re.compile(r'\bmano\b', re.IGNORECASE): 'amigo',
    re.compile(r'\bmina\b', re.IGNORECASE): 'garota',
    re.compile(r'\bgalera\b', re.IGNORECASE): 'pessoal',
    re.compile(r'\bdaí\b', re.IGNORECASE): 'então',
    re.compile(r'\bfalar\b(?=.*\bcom\b)', re.IGNORECASE): 'conversar com',
}

# Conjugação pessoal errada
VERB = {
    re.compile(r'\b(é as coisas)\b', re.IGNORECASE): 'são as coisas',
    re.compile(r'\bfique aí com\b', re.IGNORECASE): 'fiquem aí com',
    re.compile(r'\btu vai\b', re.IGNORECASE): 'tu vais',
    re.compile(r'\btu tem\b', re.IGNORECASE): 'tu tens',
    re.compile(r'\btu pode\b', re.IGNORECASE): 'tu podes',
    re.compile(r'\btu deve\b', re.IGNORECASE): 'tu deves',
    re.compile(r'\bvocê foram\b', re.IGNORECASE): 'vocês foram',
    re.compile(r'\beles foi\b', re.IGNORECASE): 'eles foram',
    re.compile(r'\belas foi\b', re.IGNORECASE): 'elas foram',
    re.compile(r'\bnós fomos\b', re.IGNORECASE): 'nós fomos',
    re.compile(r'\bnós éramos\b', re.IGNORECASE): 'nós éramos',
    re.compile(r'\beu estava\b', re.IGNORECASE): 'eu estava',
    re.compile(r'\btu estava\b', re.IGNORECASE): 'tu estavas',
    re.compile(r'\bele estava\b', re.IGNORECASE): 'ele estava',
    re.compile(r'\bnós estava\b', re.IGNORECASE): 'nós estávamos',
    re.compile(r'\bvocês estava\b', re.IGNORECASE): 'vocês estavam',
    re.compile(r'\beles estava\b', re.IGNORECASE): 'eles estavam',
    re.compile(r'\bfazer eles\b', re.IGNORECASE): 'fazê-los',
    re.compile(r'\bver eles\b', re.IGNORECASE): 'vê-los',
    re.compile(r'\bencontrar eles\b', re.IGNORECASE): 'encontrá-los',
    re.compile(r'\bcomprar eles\b', re.IGNORECASE): 'comprá-los',
    re.compile(r'\bbeber eles\b', re.IGNORECASE): 'bebê-los',
}

# Regras especiais de acordo com as suas especificações
REGRAS_ESPECIAIS = {
    re.compile(r'\bOxitocina\b', re.IGNORECASE): 'Ocitocina',
    re.compile(r'\bUNICEF\b', re.IGNORECASE): 'Unicef',
    re.compile(r'\banálise do comportamento aplicada\b', re.IGNORECASE): 'Análise do Comportamento Aplicada',
    re.compile(r'\bLei 10\.406\b', re.IGNORECASE): 'Lei no 10.406/2002',
    re.compile(r'\bet al\.\b', re.IGNORECASE): 'et al.',
    re.compile(r'\bIn:\b', re.IGNORECASE): 'In:',
    re.compile(r'\bin memorian\b', re.IGNORECASE): 'in memoriam',
    re.compile(r'\bEpicteto\b', re.IGNORECASE): 'Epiteto',
    re.compile(r'\bNum\b', re.IGNORECASE): 'em um',
    re.compile(r'\bNuma\b', re.IGNORECASE): 'em uma',
    re.compile(r'\bNo momento que\b', re.IGNORECASE): 'no momento em que',
    re.compile(r'\bEstatuto do Idoso\b', re.IGNORECASE): 'Estatuto da Pessoa Idosa',
    re.compile(r'\bonline\b', re.IGNORECASE): 'on-line',
    re.compile(r'\bNos dias atuais\b', re.IGNORECASE): 'atualmente',
    re.compile(r'\bTratam-se\b', re.IGNORECASE): 'Trata-se',
    re.compile(r'\bSéculo XXI\b', re.IGNORECASE): 'século XXI',
    re.compile(r'\bdécada de 10\b', re.IGNORECASE): 'década de 1910',
    re.compile(r'\bdécada de 20\b', re.IGNORECASE): 'década de 1920',
    re.compile(r'\bdécada de 30\b', re.IGNORECASE): 'década de 1930',
    re.compile(r'\bdécada de 40\b', re.IGNORECASE): 'década de 1940',
    re.compile(r'\bdécada de 50\b', re.IGNORECASE): 'década de 1950',
    re.compile(r'\bdécada de 60\b', re.IGNORECASE): 'década de 1960',
    re.compile(r'\bdécada de 70\b', re.IGNORECASE): 'década de 1970',
    re.compile(r'\bdécada de 80\b', re.IGNORECASE): 'década de 1980',
    re.compile(r'\bdécada de 90\b', re.IGNORECASE): 'década de 1990',
}

# Cacoetes de linguagem (mantidos do original)
CACOETES = {
    re.compile(r'\b(Aí|Ai|Aih) eu fui\b', re.IGNORECASE): 'Então eu fui',
    re.compile(r'\b(Aí|Ai|Aih) eu disse\b', re.IGNORECASE): 'Então eu disse',
    re.compile(r'\btipo assim\b', re.IGNORECASE): 'por exemplo',
    re.compile(r'\b(okay|ok)\b', re.IGNORECASE): 'certo',
    re.compile(r'\b(entendeu\?|sacou\?)\b', re.IGNORECASE): '',
    re.compile(r'\b(na real|sabe)\b', re.IGNORECASE): '',
    re.compile(r'\b(então assim)\b', re.IGNORECASE): 'portanto',
    re.compile(r'\b(tá bom|beleza)\b', re.IGNORECASE): 'certo',
    re.compile(r'\bmega\b', re.IGNORECASE): 'muito',
    re.compile(r'\bsuper\b', re.IGNORECASE): 'muito',
    re.compile(r'\bhiper\b', re.IGNORECASE): 'muito',
    re.compile(r'\bultra\b', re.IGNORECASE): 'muito',
    re.compile(r'\btop\b', re.IGNORECASE): 'excelente',
    re.compile(r'\bmassa\b', re.IGNORECASE): 'bom',
    re.compile(r'\blegal\b', re.IGNORECASE): 'bom',
    re.compile(r'\bshow\b', re.IGNORECASE): 'excelente',
}

# Gírias (mantidas do original)
GIRIAS = {
    re.compile(r'\bnois\b', re.IGNORECASE): 'nós',
    re.compile(r'\btamo\b', re.IGNORECASE): 'estamos',
    re.compile(r'\bta\b', re.IGNORECASE): 'está',
    re.compile(r'\b(vc|vcê)\b', re.IGNORECASE): 'você',
    re.compile(r'\bpra\b', re.IGNORECASE): 'para',
    re.compile(r'\bpros?\b', re.IGNORECASE): 'para os',
    re.compile(r'\bnéh?\b', re.IGNORECASE): '',
    re.compile(r'\bblz\b', re.IGNORECASE): 'certo',
    re.compile(r'\b(kkk+|rsrs)\b', re.IGNORECASE): '',
    re.compile(r'\bpq\b', re.IGNORECASE): 'porque',
    re.compile(r'\b(q|qui)\b', re.IGNORECASE): 'que',
    re.compile(r'\bmano\b', re.IGNORECASE): 'amigo',
    re.compile(r'\btbm\b', re.IGNORECASE): 'também',
}

# Configura o corretor ortográfico
spell = SpellChecker(language='pt')

def aplicar_correcoes_com_diff(doc):
    """
    NOVA VERSÃO DA FUNÇÃO PRINCIPAL.
    Esta função aplica as correções e reconstrói o parágrafo mostrando as
    diferenças (diffs) com formatação especial (riscado/vermelho, negrito/verde).
    Isso substitui a simples marcação em amarelo e preserva a formatação original
    nas partes do texto que não foram alteradas.
    """
    report = []
    # Combina todos os dicionários de regras
    todas_regras = {
        **ERRORS, **CONCORD, **HOM, **PREP, **STYLE, 
        **INFORMAL, **VERB, **REGRAS_ESPECIAIS, 
        **CACOETES, **GIRIAS
    }
    dmp = dmp_module.diff_match_patch()

    for p in doc.paragraphs:
        if not p.text.strip():
            continue

        texto_original = p.text
        texto_corrigido = texto_original

        # Etapa 1: Aplica todas as regras de substituição para gerar o texto final
        for pattern, replacement in todas_regras.items():
            texto_corrigido, count = pattern.subn(replacement, texto_corrigido)
            if count > 0:
                # O relatório é simplificado pois o diff mostrará as mudanças
                pass

        # Etapa 2: Aplica a correção ortográfica no texto já corrigido
        palavras = re.findall(r'\b\w+\b', texto_corrigido)
        palavras_erradas = spell.unknown(palavras)
        texto_final = texto_corrigido
        for palavra in palavras_erradas:
            if palavra.isupper() or palavra[0].isupper():
                continue
            correcao = spell.correction(palavra)
            if correcao and correcao != palavra:
                texto_final = re.sub(r'\b' + re.escape(palavra) + r'\b', correcao, texto_final)

        # Etapa 3: Limpa espaços duplos e múltiplos
        texto_final = re.sub(r'\s{2,}', ' ', texto_final).strip()
        
        # Etapa 4: Compara o texto original com o final e formata as diferenças
        if texto_final != texto_original:
            # Gera a lista de diferenças
            diffs = dmp.diff_main(texto_original, texto_final)
            dmp.diff_cleanupSemantic(diffs) # Otimiza os diffs para serem mais legíveis

            # Limpa o parágrafo original para reescrevê-lo com a formatação de diff
            p.clear()

            for op, text in diffs:
                run = p.add_run(text)
                # op -1: Deleção. Formata como vermelho e riscado.
                if op == dmp.DIFF_DELETE:
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    report.append({'original': text, 'corrigido': '', 'tipo': 'Remoção'})
                # op 1: Inserção. Formata como verde e negrito.
                elif op == dmp.DIFF_INSERT:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    report.append({'original': '', 'corrigido': text, 'tipo': 'Adição'})
                # op 0: Igual. Nenhuma formatação extra.
                elif op == dmp.DIFF_EQUAL:
                    pass

    return doc, report


@app.route('/revisar-documento', methods=['POST'])
def revisar_documento():
    if 'file' not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    file = request.files['file']
    if not file or file.filename == '':
        return jsonify({"error": "Nome de arquivo vazio ou inválido"}), 400

    if file and file.filename.endswith('.docx'):
        try:
            doc = docx.Document(io.BytesIO(file.read()))
            # Chama a NOVA função que aplica o diff visual
            doc_corrigido, relatorio = aplicar_correcoes_com_diff(doc)

            buffer = io.BytesIO()
            doc_corrigido.save(buffer)
            buffer.seek(0)

            response = send_file(
                buffer,
                as_attachment=True,
                download_name=f"revisado_{file.filename}",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            response.headers['X-Corrections-Report'] = json.dumps(relatorio)
            response.headers['Access-Control-Expose-Headers'] = 'X-Corrections-Report'

            return response

        except Exception as e:
            return jsonify({"error": f"Erro interno ao processar o arquivo: {str(e)}"}), 500

    return jsonify({"error": "Formato de arquivo inválido. Por favor, envie um .docx"}), 400


if __name__ == '__main__':
    app.run(debug=True, port=5000)

